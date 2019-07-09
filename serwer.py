import anvil.server
import anvil.media
import matplotlib.pyplot as plt
import math
import csv
import os
import datetime
import time
import IPython
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

#twozenie wykresów 
@anvil.server.callable
def graphing1(tablica1, wariant):
	data_x = []
	data_y = []
	for arrey in tablica1:
		data_x.append(arrey[1])
		data_y.append(arrey[8])
	plt.plot(data_x, data_y, 'bs')
	plt.ylabel('Temperatura wylotowa [k]')
	plt.xlabel('Strumień masowy [kg/s]')
	plt.title('Wykres zależności temperatury wylotowej od natężenia strumienia masowego.', fontsize=10)
	plt.savefig('wykresy/%s.png' % wariant)
	plt.close('all')
	media_object = anvil.media.from_file('wykresy/%s.png' % wariant)
	return media_object
   
@anvil.server.callable
def calculation(input_data):
    # obliczenie danych dla pierwszego zadania
    tablica1 = []
    for Re in range(100,100001,100):
      tablica1.append(wyciep_projekt_2_zad1(input_data, Re))    
    # zapisuje tablice 1 w csv
    csv_save(input_data[14], tablica1)
    # obliczenie danych do drugiego zadania
    tablica2 = []
    tablica2.append(wyciep_projekt_2_zad2(input_data))
    return (tablica1, tablica2)
	
@anvil.server.callable
def wyciep_projekt_2_zad1(input_data ,re):
	# metoda ta jest szablonem który zaczyna projekt do wyciepu zad1 lam
	# bierze adres liste z wartosicami 
	# zwraca arrey [re, m, Q, alfap, Tscian, Tpsr, alfaw, Qk, Tout_olb] 
	# czyli wszystko co trzeba do tabeli ktora wstawiamy
    #wczytanie danych 
    Tw = input_data[0]#temperatura wlotowa
    Tp = input_data[1]#temperatura powietrza
    L = input_data[2]#długość rurociagu
    dw = input_data[3]#srednica wewnetrzna 
    dz = dw + 0.0008 #srednica zewnetrza 
    Tout_zal = float(360.0) #zalozona temepratura wylotowa
    g = float(9.81)
    
    #laminarny z re = 0 do re = 2000
    while True:
        Tsr = float((Tw + Tout_zal)/2)
        #obliczenie danych fizyko chemicznych 1
        ro = 0.00000000291666678309*(Tsr**6) - 0.00000547595854757804*(Tsr**5) + 0.004279770733921*(Tsr**4) - 1.78226490008235*(Tsr**3) + 417.089905377841*(Tsr**2) - 52006.6149207153*(Tsr) + 2700228.1586917
        Cp = float((-0.00000007499999932215)*(Tsr**6) + 0.00014141749879883500*(Tsr**5) - 0.111007828583009*(Tsr**4) + 46.4322963562333*(Tsr**3) - 10914.9954438363*(Tsr**2) + 1367213.38891221*(Tsr) - 71288364.0207992)
        mi = 0.00000000000007916667*(Tsr**6) - 0.00000000014970458337*(Tsr**5) + 0.00000011787731424344*(Tsr**4) - 0.0000494720950523304*(Tsr**3) + 0.0116729112319278*(Tsr**2) - 1.46826444758532*(Tsr) + 76.9281402925932
        Pr = 0.00000000037500000018*(Tsr**6) - 0.00000071333750029697*(Tsr**5) + 0.000565200085098775*(Tsr**4) - 0.238782016317006*(Tsr**3) + 56.7371074551626*(Tsr**2) - 7190.2084859*(Tsr) + 379757.610772943
        la = (-0.00000000000138888879)*(Tsr**6) + 0.00000000290124981441*(Tsr**5) - 0.00000250312671648917*(Tsr**4) + 0.00114314816587917*(Tsr**3) - 0.291743840739458*(Tsr**2) + 39.4850991524641*(Tsr) - 2215.18263706133
        mis = 0.00000000000007916667*(Tsr**6) - 0.00000000014970458337*(Tsr**5) + 0.00000011787731424344*(Tsr**4) - 0.0000494720950523304*(Tsr**3) + 0.0116729112319278*(Tsr**2) - 1.46826444758532*(Tsr) + 76.9281402925932
        u = float((re*mi)/(dw*ro))
        m = float(((math.pi*(dw**2))/4)*ro*u)
        Q = m*Cp*(Tw-Tout_zal)
        Prs = mis/(7841*14.16e-6)
        if re <= 2000:
            Nu = 1.86*((Pr*re*dw/L)**(1/3))*((mi/mis)**0.14)
        elif re>2000 and re <= 10000:
            if re <=2200:
                k0 = 2.2
            elif re <=2300:
                k0=2.3
            elif re <=2500:
                k0=4.9   
            elif re <=3000:
                k0=7.5
            elif re <=3500:
                k0=10
            elif re <=4000:
                k0=12.2   
            elif re <=5000:
                k0=16.5
            elif re <=6000:
                k0=20
            elif re <=7000:
                k0=24   
            elif re <=8000:
                k0=27
            elif re <=9000:
                k0=30
            elif re <=10000:
                k0=33
            Nu = k0 * (Pr**0.43) * ((Pr/Prs)**0.25)
        elif re > 10000 and re <= 100000:
            Nu = 0.08*((re)**0.8)*((Pr)**0.33)*((mi/mis)**0.14)
        alfaw = Nu*la/dz
        Tscian = Tsr - (Q/(alfaw*math.pi*dw*L))
        Tsczew = Tscian - (Q/L)*(math.log(dz/dw)/(2*math.pi*la))
        Tpsr = (Tsczew + Tp)/2
        #obliczenie danych fizyko chemicznych 2
        lap = 0.00000000000208333333714955*(Tp**6) - 0.00000000377687500697432*(Tp**5) + 0.00000285021737510634*(Tp**4) - 0.00114606219959937*(Tp**3) + 0.258970652992879*(Tp**2) - 31.1805545881856*(Tp) + 1562.80895539821
        nu = (-0.000000000000000069444440790223)*(Tp**6) + 0.000000000000125895827114741*(Tp**5) - 0.0000000000949794634341394*(Tp**4) + 0.00000003816783264417*(Tp**3) - 0.00000861641891404772*(Tp**2) + 0.00103613194866617*(Tp) - 0.051845870774733
        Prp = (-0.00000000000694444431)*(Tp**6) + 0.00000001250624977570*(Tp**5) - 0.00000937580079287315*(Tp**4) + 0.00374537760424578*(Tp**3) - 0.840846459899592*(Tp**2) + 100.589084989333*(Tp) - 5008.73838670239
        las = 50.3
        Gr = (g*(dz**3)*(Tsczew - Tp))/(Tp*(nu**2))
        #wylicznie stałych c oraz n w zalezności od Gr i prp
        if Gr*Prp <= 5e-02:
            C=1.18
            n=0.125
        elif Gr*Prp > 5e02 and Gr*Prp <= 2e07:
            C=0.54
            n=0.25
        elif Gr*Prp > 2e07:
            C=0.135
            n=1/3
        Nu = C*((Gr*Prp)**n)
        alfap = (Nu*lap)/dz
        R = float(1/(math.pi*alfaw*dw)) + (1/(2*math.pi*las))*math.log(dz/dw) + (1/(math.pi*alfap*dz))
        K = float(1/R)
        Tout_olb = Tp + (Tw-Tp)/math.exp((K*L)/(m*Cp))
        delT = ((Tw-Tp)-(Tout_olb-Tp))/math.log((Tw-Tp)/(Tout_olb-Tp))
        Qk=(L*delT)/R
        if round(Tout_olb,4) == round(Tout_zal,4):
            return [re, m, Q, alfap, Tscian, Tpsr, alfaw, Qk, Tout_olb]
        else: 
            #print(Tout_olb)
            Tout_zal = round(Tout_olb,10)

@anvil.server.callable
def wyciep_projekt_2_zad2(input_data):
    #wczytanie danych 
    rodzaj_wymienika = input_data[5]#liczby 1 dla kwadratowego 2 dla okraglego
    T1 = input_data[6]#temperatura 1
    T2 = input_data[7]#temperatura 2
    m = input_data[8]#strumien masy
    rodzaj_przeplywu = input_data[9]#litery p albo w 
    t1 = input_data[10]
    t2 = input_data[11]
    mz = (m*1000)/(3600*24)
    Tsrz = (T1+T2)/2
    Tsrg = (t1+t2)/2
    ug = 1.8
    uz = ug 
    #obliczenie danych fizyko chemicznych 1
    Cpz = (-0.00000007499999932215)*(Tsrz**6) + 0.00014141749879883500*(Tsrz**5) - 0.111007828583009*(Tsrz**4) + 46.4322963562333*(Tsrz**3) - 10914.9954438363*(Tsrz**2) + 1367213.38891221*(Tsrz) - 71288364.0207992
    Cpg = (-0.00000007499999932215)*(Tsrg**6) + 0.00014141749879883500*(Tsrg**5) - 0.111007828583009*(Tsrg**4) + 46.4322963562333*(Tsrg**3) - 10914.9954438363*(Tsrg**2) + 1367213.38891221*(Tsrg) - 71288364.0207992
    rog = 0.00000000291666678309*(Tsrg**6) - 0.00000547595854757804*(Tsrg**5) + 0.004279770733921*(Tsrg**4) - 1.78226490008235*(Tsrg**3) + 417.089905377841*(Tsrg**2) - 52006.6149207153*(Tsrg) + 2700228.1586917
    roz = 0.00000000291666678309*(Tsrz**6) - 0.00000547595854757804*(Tsrz**5) + 0.004279770733921*(Tsrz**4) - 1.78226490008235*(Tsrz**3) + 417.089905377841*(Tsrz**2) - 52006.6149207153*(Tsrz) + 2700228.1586917
    mi = 50.3
    miz = 0.00000000000007916667*(Tsrz**6) - 0.00000000014970458337*(Tsrz**5) + 0.00000011787731424344*(Tsrz**4) - 0.0000494720950523304*(Tsrz**3) + 0.0116729112319278*(Tsrz**2) - 1.46826444758532*(Tsrz) + 76.9281402925932
    mig = 0.00000000000007916667*(Tsrg**6) - 0.00000000014970458337*(Tsrg**5) + 0.00000011787731424344*(Tsrg**4) - 0.0000494720950523304*(Tsrg**3) + 0.0116729112319278*(Tsrg**2) - 1.46826444758532*(Tsrg) + 76.9281402925932
    Prg = 0.00000000037500000018*(Tsrg**6) - 0.00000071333750029697*(Tsrg**5) + 0.000565200085098775*(Tsrg**4) - 0.238782016317006*(Tsrg**3) + 56.7371074551626*(Tsrg**2) - 7190.2084859*(Tsrg) + 379757.610772943
    Prz = 0.00000000037500000018*(Tsrz**6) - 0.00000071333750029697*(Tsrz**5) + 0.000565200085098775*(Tsrz**4) - 0.238782016317006*(Tsrz**3) + 56.7371074551626*(Tsrz**2) - 7190.2084859*(Tsrz) + 379757.610772943
    Q = mz * Cpz * (T2-T1)
    mg = Q/(Cpg*(t1-t2))
    dw = math.sqrt((4*mz)/(math.pi*roz*0.9*uz))
    dz = dw + 0.005 
    Dw = math.sqrt(((4*mg)/(math.pi*rog*0.9*ug))+(dz**2))
    if rodzaj_przeplywu == "w":
        deltaTlog = ((t1-T1)-(t2-T2))/(math.log((t1-T1)/(t2-T2)))
    elif rodzaj_przeplywu == "p":
        deltaTlog = ((t1-T2)-(t2-T1))/(math.log((t1-T2)/(t2-T1)))
    else:
        return "zle wprowadzony rodzaj przeplywu"
    #wstepne zalozenia temperatury 
    Tsczi_zal=400
    Tscgo_zal=400
    while True:
        while True:
            #obliczenie danych fizyko chemicznych 2
            Prgsci = 0.00000000037500000018*(Tscgo_zal**6) - 0.00000071333750029697*(Tscgo_zal**5) + 0.000565200085098775*(Tscgo_zal**4) - 0.238782016317006*(Tscgo_zal**3) + 56.7371074551626*(Tscgo_zal**2) - 7190.2084859*(Tscgo_zal) + 379757.610772943
            Przsci = 0.00000000037500000018*(Tsczi_zal**6) - 0.00000071333750029697*(Tsczi_zal**5) + 0.000565200085098775*(Tsczi_zal**4) - 0.238782016317006*(Tsczi_zal**3) + 56.7371074551626*(Tsczi_zal**2) - 7190.2084859*(Tsczi_zal) + 379757.610772943
            la = 50.3
            laz = (-0.00000000000138888879)*(Tsrz**6) + 0.00000000290124981441*(Tsrz**5) - 0.00000250312671648917*(Tsrz**4) + 0.00114314816587917*(Tsrz**3) - 0.291743840739458*(Tsrz**2) + 39.4850991524641*(Tsrz) - 2215.18263706133
            lag = (-0.00000000000138888879)*(Tsrg**6) + 0.00000000290124981441*(Tsrg**5) - 0.00000250312671648917*(Tsrg**4) + 0.00114314816587917*(Tsrg**3) - 0.291743840739458*(Tsrg**2) + 39.4850991524641*(Tsrg) - 2215.18263706133
            Rez = (0.9 * uz * dw * roz)/miz
            Reg = (0.9 * ug * (Dw-dz) * rog)/mig
            mizz = (Prz/Przsci)**0.25
            migg = (Prg/Prgsci)**0.25
            Nuz = 0.021 * (Rez**0.8) * (Prz**0.43) * mizz
            Nug = 0.02 * (Reg**0.8) * (Prg**0.43) * migg
            alfaz = (Nuz*laz)/dw
            alfag = (Nug*lag)/(Dw-dz)
            R = (1/(math.pi*alfaz*dw)) + (1/(2*math.pi*la))*math.log(dz/dw) + (1/(math.pi*alfag*dz))
            L = (Q*R)/(deltaTlog)
            Tsczi_obl= Tsrz - (Q/(alfaz*math.pi*dw*L))
            Tscgo_obl= Tsrg - (Q/(alfag*math.pi*dz*L))
            if round(Tsczi_obl,8) == round(Tsczi_zal,8):
                break
            else:
                Tsczi_zal = Tsczi_obl

        if round(Tscgo_zal,8) == round(Tscgo_obl,8):
            return [T1,T2,m,t1,t2,ug,uz,Q,mg,dw,dz,Dw,deltaTlog,Tsczi_obl,Tscgo_obl,Rez,Reg,Nuz,Nug,alfag,alfaz,R,L]
        else:
            Tscgo_zal = Tscgo_obl

@anvil.server.callable
def csv_save(number_of_wariant, tablica):
	#klasa zapisujaca plik csv
    with open(os.path.expanduser("C:/Users/KORNEL/Dropbox/projekty_ichip/wyciep python/csv_files/"+str(number_of_wariant)+".csv"), 'a',newline='') as csvfile:
        spamwriter = csv.writer(csvfile, delimiter=';')
        for iteration in tablica:
            spamwriter.writerow(iteration)

@anvil.server.callable
def get_csv(wariant):
	media_object = anvil.media.from_file('csv_files/%s.csv' % wariant)
	return media_object

#klasa towzaca worda do projektu wyciep 2
def docx_save(imie, nazwisko, numer_warantu, tablica1, tablica2):
    #zadanie 1
    document = Document()
    document.add_heading('%d %s %s' % (datetime.datetime.now().year,imie,nazwisko),0)
    document.add_heading('WYCIEP projekt 2',0)
    document.add_page_break()
    document.add_heading('ZADANIE 1 ',1)
    document.add_paragraph('Przez rurociąg wykonany ze stali węglowej (0,15% C) o średnicy wewnętrznej dw [m] i zewnętrznej dz=dw+0,008 [m] oraz długości L [m] przepływa woda [ton/dobę] o temperaturze wlotowej Tw [K]. Rurociąg jest umieszczony w hali produkcyjnej, w której panuje stała temperatura powietrza równa Tp [K]. Określić straty cieplne i temperaturę wylotową cieczy w zależności od natężenia masowego przepływu w zakresie liczby Reynoldsa Re=100-105. Rozwiązanie podać w formie wykresu. ')
    document.add_heading('Analiza problemu',1)
    document.add_paragraph('Problem przedstawiony w zadaniu polega na wyznaczeniu temperatury wylotowej, którą mamy uzależnić od natężenia masowego przepływu w zakresie liczby Reynoldsa Re=100- 105. W zadaniu mamy powiedziane, że przez rurociąg przepływa woda, więc mamy do czynienia z mechanizmem konwekcyjnym. Wewnątrz rurociągu występuje konwekcja wymuszona, ponieważ występuje ruch płynu, natomiast w jego otoczeniu- swobodna, gdyż powietrze tylko otacza rurociąg i jest to ruch samoistny.')
    document.add_paragraph('Ciepło jakie niesie ze sobą ciecz przenika przez ścianki (zakładam ze rurociąg jest ze stali 0.15% C- [2] Gogół tablica 4.1 str.53) – na początku zachodzi wnikanie od płynu do ścianki, później przewodzenie wewnątrz ścianki a na końcu wnikanie ze ścianki do powietrza.')
    document.add_paragraph('W zależności od konwekcji liczba Nusselta jest funkcją liczb: Reynoldsa i Prandtla dla wymuszonej, Prandtla i Grashofa dla swobodnej. Liczbą Nusselta nazywamy stosunek szybkości wnikania ciepła do szybkości przewodzenia ciepła. Liczba Reynoldsa to siły bezwładności przez siły tarcia. Liczba Prandtla to siła molekularnego przenoszenia pędu przez szybkość przewodzenia ciepła. Liczba Grashofa to siła wyporu powstała na skutek różnicy temperatur przez siły tarcia. Są to liczby kryterialne.')
    document.add_page_break()
    document.add_heading('Koncepcja rozwiązania',1)
    document.add_paragraph('Własności fizykochemiczne cieczy są zależne od temperatury', style='ListBullet')
    document.add_paragraph('Wewnątrz rurociągu występuje konwekcja wymuszona', style='ListBullet')
    document.add_paragraph('Pole temperatur jest ustalone', style='ListBullet')
    document.add_paragraph('Temperatura ścianki wielkością stalą na całej długości rurociągu', style='ListBullet')
    document.add_heading('Zakładam temperaturę wylotową, a następnie obliczam temperaturę średnią :',3)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Tsr=(Tw-Twylot)/2'
    hdr_cells[2].text = ''
    document.add_paragraph('Tw – temperatura wlotowa wody')
    document.add_paragraph('Twylot – założona temperatura wylotowa wody')
    document.add_paragraph('')
    document.add_paragraph('Następnie korzystając z tabel zawartych w książce Wiesława Gogóła „Wymiana ciepła – tablice i wykresy” interpoluję potrzebne wielkości fizykochemiczne w Excelu.')
    document.add_paragraph('ρ  – gęstość wody[kg/m3]')
    document.add_paragraph('μ – lepkość dynamiczna wody[Pa*s]')
    document.add_paragraph('Cp– ciepło właściwe wody [J/(kg*K)]')
    document.add_paragraph('λ– współczynnik przewodzenia ciepła wody')
    document.add_heading('Obliczam prędkość przepływu wody:',3)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Re=(u*dw*ρ)/μ => u=(Re*μ)/(dw*ρ)'
    hdr_cells[2].text = '' 
    document.add_paragraph('Re- Liczba Reynoldsa w zakresie 100<Re<105')
    document.add_paragraph('μ– lepkość dynamiczna wody [Pa*s] ')
    document.add_paragraph('ρ – gęstość wody [kg/m3]')
    document.add_paragraph('dw– średnica wewnętrzna rury [m]')
    document.add_page_break()
    document.add_heading('Obliczam natężenie masowe przepływu:',3)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'm[kg/s]=(u*ρ*π*dw^2)/4'
    hdr_cells[2].text = '' 
    document.add_heading('Z bilansu entalpowego obliczam natężenie przepływu:',3)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Q=m*Cp*(Tw-Twylot)'
    hdr_cells[2].text = '' 
    document.add_paragraph('m [kg/s] – natężenie masowe przepływu')
    document.add_paragraph('Cp [J/(kg*K)] – ciepło właściwe wody')
    document.add_paragraph('Tw [K] – temperatura wlotowa wody')
    document.add_paragraph('Twylot  [K] – założona temperatura wylotowa wody')
    document.add_paragraph('Następnie interpoluję liczbę Prandtla oraz zakładam temperaturę ścianki rurociągu.')
    document.add_heading('Następnie korzystając z odpowiedniej korelacji dla określonego przepływu obliczam liczbę Nusselta:',3)
    document.add_paragraph('a) przepływ laminarny (Re<2000). Długość rury wywiera duży wpływ – ze wzrostem długości rury średnia wartość α zmienia się (Hobler str.195):')
    document.add_paragraph('                 Nu = 1.86*((Pr*Re*dw/L)^(1/3))*((μ/μs)^0.14)')
    document.add_paragraph('b) przepływ przejściowy (2000<Re<10000) Gogół tablica 6.1 str.102:')
    document.add_paragraph('                 Nu = K0 * Pr^0,43 * (Pr*λsc / Cpsc*μsc)^0,25')
    document.add_paragraph('')
    table = document.add_table(rows=2, cols=13)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Re*e-3'
    hdr_cells[1].text = '2,2'
    hdr_cells[2].text = '2,3'
    hdr_cells[3].text = '2,5' 
    hdr_cells[4].text = '3,0'
    hdr_cells[5].text = '3,5'
    hdr_cells[6].text = '4'
    hdr_cells[7].text = '5'
    hdr_cells[8].text = '6'
    hdr_cells[9].text = '7'
    hdr_cells[10].text = '8'
    hdr_cells[11].text = '9'
    hdr_cells[12].text = '10'
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = 'K0'
    hdr_cells[1].text = '2,2'
    hdr_cells[2].text = '3,6' 
    hdr_cells[3].text = '4,9'
    hdr_cells[4].text = '7,5'
    hdr_cells[5].text = '10'
    hdr_cells[6].text = '12,2'
    hdr_cells[7].text = '16,5'
    hdr_cells[8].text = '20'
    hdr_cells[9].text = '24'
    hdr_cells[10].text = '27'
    hdr_cells[11].text = '30'
    hdr_cells[12].text = '33'
    document.add_paragraph('')
    document.add_paragraph('c) przepływ burzliwy (Re>10000 i 0,7<Pr<16700):')
    document.add_paragraph('                 Nu = 0,023 * Re^0,8 * Pr^0,33 * (μ/μsc)^0,14')
    document.add_heading('Następnie  znając  liczbę  Nusselta obliczam współczynnik wnikania ciepła αw. ',3)
    document.add_paragraph('                 Nu = (αw*dw)/λ  =>   αw=(Nu*λ)/dw')   
    document.add_paragraph('Nu – liczba Nusselta dla przepływu')
    document.add_paragraph('λ [W/(m*K)] – współczynnik przewodzenia ciepła wody')
    document.add_paragraph('dw [m] – średnica wewnętrzna rury')
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_heading('Obliczam właściwą temperaturę ścianki z bilansu cieplnego:',3)
    document.add_paragraph('          Q = αw * π * dw * L * (Tsr - Tścianki) => Tścianki = Tsr - Q/(αw * π * dw * L)')
    document.add_paragraph('Q [W] – strumień ciepła')
    document.add_paragraph('dw [m] – średnica wewnętrzna rury')
    document.add_paragraph('L  [m] – długość rury')
    document.add_paragraph('Tśr [K] – średnia temperatura wody w rurze')
    document.add_paragraph('Jeżeli wartość założona nie zgadza się z obliczoną podstawiam znowu temperaturę i obliczam do momentu aż założona i wyliczona będą się zgadzały ')
    document.add_paragraph('Obliczam temperaturę zewnętrzną rury z bilansu cieplnego:')
    document.add_paragraph('        Q/L=k*ΔT=ln(dz/dw)*(Tścianki - Tść.zew)/(2*π*λstali)')
    document.add_paragraph('        Tść.zew=Tścianki - Q/L * (ln(dz/dw)/(2*π*λstali))')
    document.add_paragraph('λstali  [W/(m*K)] – współczynnik przewodzenia stali (0,15% C)')
    document.add_paragraph('dz [m] – średnica zewnętrzna rury')
    document.add_paragraph('dw [m] – średnica wewnętrzna rury')
    document.add_paragraph('Tścianki [K] – temperatura wewnętrzna ścianki')
    document.add_paragraph('Q [W] – strumień ciepła')
    document.add_paragraph('L– długość rury')
    document.add_heading('Obliczam temperaturę warstwy przyściennej jako średnią powietrza i odczytuję dane fizykochemiczne:',3)
    document.add_paragraph('        Tpśr = (Tść.zew + T)/2')
    document.add_paragraph('ρ [kg/m3] – gęstość powietrza')
    document.add_paragraph('μ [Pa*s] – lepkość dynamiczna powietrza')
    document.add_paragraph('Cp [J/(kg*K)] – ciepło właściwe powietrza')
    document.add_paragraph('λ [W/(m*K)] – współczynnik przewodzenia ciepła powietrza')
    document.add_paragraph('Pr [-] - liczba Prandtla')
    document.add_heading('Obliczam liczbę Grashofa:',3)
    document.add_paragraph('        Gr = (g * dz^3 / v^2) * β * ΔT = (g * dz^3 / (μ/ρ)^2) * (Tść.zew -Tp)/Tp')
    document.add_paragraph('β - współczynnik rozszerzalności termicznej')
    document.add_paragraph('g [m/s] – przyspieszenie ziemskie')
    document.add_paragraph('dz  [m] – średnica zewnętrzna rury')
    document.add_paragraph('ρ  [kg/m3] – gęstość powietrza')
    document.add_paragraph('μ [Pa*s] – lepkość dynamiczna powietrza')
    document.add_paragraph('Tść.zew. [K] – temperatura zewnętrznej ścianki rury ')
    document.add_paragraph('Tp [K]  - temperatura powietrza')
    document.add_heading('Korzystając z odpowiedniej korelacji dla określonego przepływu obliczam liczbę Nusselta dla konwekcji swobodnej (wzór Michiejewa):',3)
    document.add_paragraph('        Nu = C * (Gr * Pr)^n')
    document.add_paragraph('Stałe C i n zależą od:')
    table = document.add_table(rows=4, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Gr * Pr'
    hdr_cells[1].text = 'C'
    hdr_cells[2].text = 'n'
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = '10^-3 do 5*10^2'
    hdr_cells[1].text = '1,18'
    hdr_cells[2].text = '1/8'
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = '5*10^2 do 2*10^7'
    hdr_cells[1].text = '0,54'
    hdr_cells[2].text = '1/4'
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '2*10^7 do 10^13'
    hdr_cells[1].text = '0,135'
    hdr_cells[2].text = '1/3'
    document.add_heading('Znając liczbę Nusselta dla powietrza obliczam współczynnik wnikania ciepła od zewnętrznej ścianki rury do powietrza αp.',3)
    document.add_paragraph('        Nu = (αp *dz)/λpow   =>  αp = (Nu * λpow)/dz')
    document.add_heading('Obliczam współczynnik oporu R',3)
    document.add_paragraph('      R = 1/(αw * dw * π) + 1/(2*π*λstali)*ln(dz/dw) + 1/(αp * dz * π)')
    document.add_paragraph('αw [W/(m2*K)] – współczynnik wnikania ciepła po stronie wody')
    document.add_paragraph('αp  [W/(m2*K)] - współczynnik wnikania ciepła od zewnętrznej ścianki rury do powietrza ')
    document.add_paragraph('λstali [W/(m*K)] – współczynnik przewodzenia stali ( 0,15% C)')
    document.add_paragraph('dz [m] – średnica zewnętrzna rury')
    document.add_paragraph('dw [m] – średnica wewnętrzna rury')
    document.add_paragraph('')  
    #wukresy do worda
    print(numer_warantu, "numer wariantu")
    document.add_picture('wykresy/%s.png' %numer_warantu)
    document.save('wordy zrobione/%s.docx' %numer_warantu)

@anvil.server.callable
def get_docx(imie, nazwisko, numer_warantu, tablica1, tablica2):
	docx_save(imie, nazwisko, numer_warantu, tablica1, tablica2)
	media_object = anvil.media.from_file('wordy zrobione/%s.docx' %(numer_warantu))
	return media_object

anvil.server.connect("")
anvil.server.wait_forever()
